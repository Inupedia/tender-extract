"""
文档解析层：支持 PDF/DOCX/TXT/Markdown 输入，统一转换为 Markdown 中间格式

成熟方案核心思想：
1. 解析层：将任意格式文档转换为 Markdown + Pages 中间层
2. 保留标题层级、段落结构、表格信息
3. 为后续切块和模块路由提供统一输入
"""
import os
import re
import logging
from pathlib import Path
from typing import Optional, Dict, Any
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# 可选依赖
try:
    import pymupdf  # PyMuPDF (fitz)
    PYMUPDF_AVAILABLE = True
except ImportError:
    try:
        import fitz as pymupdf
        PYMUPDF_AVAILABLE = True
    except ImportError:
        PYMUPDF_AVAILABLE = False
        logger.info("pymupdf未安装，PDF解析将不可用。安装: pip install pymupdf")

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.info("python-docx未安装，DOCX解析将不可用。安装: pip install python-docx")

try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False
    logger.info("paddleocr未安装，OCR功能将不可用。安装: pip install paddlepaddle paddleocr")


@dataclass
class ParsedDocument:
    """解析后的文档结构"""
    content: str  # Markdown格式的文本内容
    filename: str
    original_format: str  # pdf/docx/txt/md
    total_pages: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    tables: list = field(default_factory=list)  # 提取的表格数据
    page_contents: list = field(default_factory=list)  # 按页的内容


class DocumentParser:
    """
    统一文档解析器

    支持格式：
    - PDF: 通过 PyMuPDF 解析，提取文本和表格
    - PDF (扫描件): 通过 PaddleOCR 识别文字
    - DOCX: 通过 python-docx 解析，保留标题层级
    - TXT/MD: 直接读取

    智能检测策略：
    - 逐页判断是文字页还是扫描页
    - 文字页用 PyMuPDF 直接提取（快）
    - 扫描页/图片页用 PaddleOCR 识别（准）
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md', '.markdown'}

    # 扫描页判定阈值：每页文字少于此字符数视为扫描页
    SCANNED_PAGE_THRESHOLD = 50

    def __init__(self, use_ocr: bool = True):
        """
        Args:
            use_ocr: 是否启用 OCR（对扫描页自动调用 PaddleOCR）
        """
        self.use_ocr = use_ocr and PADDLEOCR_AVAILABLE
        self._ocr_engine = None
        self._check_dependencies()

    def _check_dependencies(self):
        """检查可用的解析依赖"""
        self.available_formats = {'.txt', '.md', '.markdown'}
        if PYMUPDF_AVAILABLE:
            self.available_formats.add('.pdf')
        if PYTHON_DOCX_AVAILABLE:
            self.available_formats.update({'.docx', '.doc'})

    def _get_ocr_engine(self):
        """懒加载 OCR 引擎（首次使用时初始化，避免无需时的开销）"""
        if self._ocr_engine is None and self.use_ocr:
            logger.info("初始化 PaddleOCR 引擎...")
            self._ocr_engine = PaddleOCR(
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=False,
            )
        return self._ocr_engine

    def _is_scanned_page(self, page) -> bool:
        """判断一页是否为扫描页（无可提取文字或文字过少）"""
        text = page.get_text("text").strip()
        # 文字少于阈值 → 视为扫描页
        if len(text) < self.SCANNED_PAGE_THRESHOLD:
            return True
        # 如果页面有图片且文字很少，也可能是扫描件
        images = page.get_images(full=True)
        if images and len(text) < self.SCANNED_PAGE_THRESHOLD * 3:
            # 检查图片面积占比
            page_area = page.rect.width * page.rect.height
            for img in images:
                xref = img[0]
                try:
                    img_rect = page.get_image_bbox(xref)
                    img_area = img_rect.width * img_rect.height
                    if img_area > page_area * 0.5:  # 图片占页面50%以上
                        return True
                except Exception:
                    pass
        return False

    def _ocr_page(self, page) -> str:
        """对单页进行 OCR 识别"""
        ocr = self._get_ocr_engine()
        if ocr is None:
            return ""

        # 将页面渲染为图片
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")

        # 保存临时文件供 PaddleOCR 处理
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            tmp.write(img_bytes)
            tmp_path = tmp.name

        try:
            result = ocr.predict(tmp_path)
            lines = []
            for res in result:
                if hasattr(res, 'rec_texts'):
                    lines.extend(res.rec_texts)
                elif isinstance(res, dict) and 'rec_texts' in res:
                    lines.extend(res['rec_texts'])
                else:
                    # 兼容不同版本的输出格式
                    try:
                        for item in res:
                            if hasattr(item, 'rec_texts'):
                                lines.extend(item.rec_texts)
                            elif isinstance(item, dict) and 'text' in item:
                                lines.append(item['text'])
                    except (TypeError, AttributeError):
                        pass
            return "\n".join(lines)
        except Exception as e:
            logger.warning(f"OCR 识别失败: {e}")
            return ""
        finally:
            os.unlink(tmp_path)

    def can_parse(self, file_path: str) -> bool:
        """检查是否能解析指定文件"""
        ext = Path(file_path).suffix.lower()
        return ext in self.available_formats

    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文档，统一输出 Markdown 格式

        Args:
            file_path: 文件路径

        Returns:
            ParsedDocument 包含 Markdown 内容和元数据
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.available_formats:
            raise ValueError(
                f"不支持的文件格式: {ext}。"
                f"可用格式: {self.available_formats}。"
                f"安装 pymupdf 支持 PDF，python-docx 支持 DOCX。"
            )

        logger.info(f"解析文档: {path.name} (格式: {ext})")

        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ('.docx', '.doc'):
            return self._parse_docx(file_path)
        elif ext in ('.md', '.markdown'):
            return self._parse_markdown(file_path)
        else:
            return self._parse_text(file_path)

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """
        解析 PDF 文件

        智能策略：
        - 逐页判断是文字页还是扫描页
        - 文字页：PyMuPDF 直接提取（快速）
        - 扫描页：PaddleOCR 识别文字（准确）
        - 同时提取表格结构
        """
        doc = pymupdf.open(file_path)
        pages_content = []
        all_tables = []
        markdown_parts = []
        ocr_pages = 0

        for page_num, page in enumerate(doc):
            is_scanned = self._is_scanned_page(page)

            if is_scanned and self.use_ocr:
                # 扫描页 → OCR 识别
                ocr_pages += 1
                logger.info(f"  第{page_num+1}页: 扫描页，使用 OCR 识别")
                ocr_text = self._ocr_page(page)
                pages_content.append(ocr_text)
                markdown_parts.append(f"<!-- OCR page {page_num+1} -->\n{ocr_text}")
            else:
                # 文字页 → PyMuPDF 直接提取
                page_text = page.get_text("text")
                pages_content.append(page_text)

                # 提取表格
                tables = page.find_tables()
                for table in tables:
                    table_data = table.extract()
                    if table_data:
                        all_tables.append({
                            'page': page_num + 1,
                            'data': table_data
                        })
                        markdown_parts.append(self._table_to_markdown(table_data))

                # 使用字典模式获取带格式信息的文本块
                blocks = page.get_text("dict")["blocks"]
                page_markdown = self._blocks_to_markdown(blocks, page_num + 1)
                markdown_parts.append(page_markdown)

        doc.close()

        content = "\n\n".join(markdown_parts)
        content = self._clean_pdf_content(content)

        metadata = {
            'page_count': len(pages_content),
            'ocr_pages': ocr_pages,
            'text_pages': len(pages_content) - ocr_pages,
        }
        if ocr_pages > 0:
            logger.info(f"  OCR 统计: {ocr_pages} 页扫描件已识别")

        return ParsedDocument(
            content=content,
            filename=Path(file_path).name,
            original_format='pdf',
            total_pages=len(pages_content),
            metadata=metadata,
            tables=all_tables,
            page_contents=pages_content
        )

    def _blocks_to_markdown(self, blocks: list, page_num: int) -> str:
        """将 PDF 文本块转换为 Markdown"""
        lines = []
        prev_size = 0

        for block in blocks:
            if block["type"] != 0:  # 只处理文本块
                continue

            for line_data in block.get("lines", []):
                line_text = ""
                max_size = 0
                is_bold = False

                for span in line_data.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    line_text += text
                    size = span.get("size", 12)
                    max_size = max(max_size, size)
                    if "bold" in span.get("font", "").lower() or \
                       span.get("flags", 0) & 2**4:
                        is_bold = True

                if not line_text.strip():
                    continue

                # 基于字体大小推断标题级别
                if max_size >= 18 and is_bold:
                    lines.append(f"# {line_text}")
                elif max_size >= 15 and is_bold:
                    lines.append(f"## {line_text}")
                elif max_size >= 13 and is_bold:
                    lines.append(f"### {line_text}")
                elif is_bold and len(line_text) < 50:
                    lines.append(f"#### {line_text}")
                else:
                    lines.append(line_text)

                prev_size = max_size

        return "\n".join(lines)

    def _table_to_markdown(self, table_data: list) -> str:
        """将表格数据转换为 Markdown 表格"""
        if not table_data or len(table_data) < 1:
            return ""

        lines = []
        # 表头
        header = table_data[0]
        header_cells = [str(cell or '').strip() for cell in header]
        lines.append("| " + " | ".join(header_cells) + " |")
        lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

        # 数据行
        for row in table_data[1:]:
            cells = [str(cell or '').strip() for cell in row]
            # 补齐列数
            while len(cells) < len(header_cells):
                cells.append('')
            lines.append("| " + " | ".join(cells[:len(header_cells)]) + " |")

        return "\n".join(lines)

    def _clean_pdf_content(self, content: str) -> str:
        """清理 PDF 提取的内容"""
        # 移除过多连续空行
        content = re.sub(r'\n{4,}', '\n\n\n', content)
        # 移除页眉页脚常见模式
        content = re.sub(r'第\s*\d+\s*页\s*共\s*\d+\s*页', '', content)
        content = re.sub(r'Page\s+\d+\s+of\s+\d+', '', content, flags=re.IGNORECASE)
        # 合并被分页断开的段落
        content = re.sub(r'([^\n])\n([^\n#|])', r'\1\n\2', content)
        return content.strip()

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """
        解析 DOCX 文件

        保留标题层级、段落、表格结构
        """
        doc = DocxDocument(file_path)
        markdown_parts = []
        all_tables = []

        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':
                # 段落处理
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para:
                    md_line = self._paragraph_to_markdown(para)
                    if md_line:
                        markdown_parts.append(md_line)

            elif tag == 'tbl':
                # 表格处理
                for table in doc.tables:
                    if table._element == element:
                        table_md, table_data = self._docx_table_to_markdown(table)
                        if table_md:
                            markdown_parts.append(table_md)
                            all_tables.append({'data': table_data})
                        break

        content = "\n\n".join(markdown_parts)

        return ParsedDocument(
            content=content,
            filename=Path(file_path).name,
            original_format='docx',
            total_pages=0,
            metadata={},
            tables=all_tables
        )

    def _paragraph_to_markdown(self, para) -> Optional[str]:
        """将 DOCX 段落转换为 Markdown"""
        text = para.text.strip()
        if not text:
            return None

        style_name = (para.style.name or '').lower() if para.style else ''

        # 标题检测
        if 'heading 1' in style_name or style_name == '标题 1':
            return f"# {text}"
        elif 'heading 2' in style_name or style_name == '标题 2':
            return f"## {text}"
        elif 'heading 3' in style_name or style_name == '标题 3':
            return f"### {text}"
        elif 'heading 4' in style_name or style_name == '标题 4':
            return f"#### {text}"
        elif 'heading' in style_name or '标题' in style_name:
            # 尝试提取标题级别
            level_match = re.search(r'(\d)', style_name)
            if level_match:
                level = min(int(level_match.group(1)), 6)
                return f"{'#' * level} {text}"
            return f"### {text}"

        # 列表检测
        if style_name.startswith('list') or '列表' in style_name:
            return f"- {text}"

        return text

    def _docx_table_to_markdown(self, table) -> tuple:
        """将 DOCX 表格转换为 Markdown"""
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return None, []

        md = self._table_to_markdown(rows_data)
        return md, rows_data

    def _parse_markdown(self, file_path: str) -> ParsedDocument:
        """解析 Markdown 文件（直接读取）"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return ParsedDocument(
            content=content,
            filename=Path(file_path).name,
            original_format='md',
            total_pages=1,
            metadata={}
        )

    def _parse_text(self, file_path: str) -> ParsedDocument:
        """解析纯文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return ParsedDocument(
            content=content,
            filename=Path(file_path).name,
            original_format='txt',
            total_pages=1,
            metadata={}
        )

    @staticmethod
    def get_supported_extensions() -> set:
        """获取当前环境支持的文件扩展名"""
        extensions = {'.txt', '.md', '.markdown'}
        if PYMUPDF_AVAILABLE:
            extensions.add('.pdf')
        if PYTHON_DOCX_AVAILABLE:
            extensions.update({'.docx', '.doc'})
        return extensions
